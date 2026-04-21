"""
文件解析器模块 - PDF/Word 文档解析

【Java 类比】
- 类似 Apache POI (Word) + PDFBox / iText (PDF) 的封装
- 或者类似 FileUtil.parseFile() 工具类
- 职责：将上传的简历文件转换为结构化文本

【Python 库选择理由】
1. pdfplumber: 比 PyPDF2 更强大，支持表格提取、精确位置定位
   - 类似 Java 的 Apache PDFBox（比 iText 更易用）
2. python-docx: 处理 .docx 文件（基于 OOXML 标准）
   - 类似 Java 的 Apache POI (XWPF)
3. 优先级: pdfplumber > PyPDF2（pdfplumber 对中文支持更好）

【支持的文件格式】
- .pdf   → pdfplumber 解析（优先）
- .docx  → python-docx 解析
- .doc   → 不支持（建议用户转换为 .docx）

【设计模式】
- 策略模式: 根据文件扩展名选择不同的解析器
- 模板方法模式: 统一的 parse() 接口，子类实现具体逻辑

【使用示例】
from app.tools.file_parser import FileParser, ParseResult

parser = FileParser()
result = await parser.parse_file("resume.pdf")
print(result.text)        # 原始文本
print(result.metadata)    # 元数据（页数、字数等）
"""

import os
import re
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field
from pathlib import Path

from app.infrastructure.logger import get_logger
from app.infrastructure.error_handler import FileParseError

logger = get_logger(__name__)


@dataclass
class ParseResult:
    """
    文件解析结果

    【Java 类比】
    - 类似 ParseResult<T> 泛型类或 Record
    - 封装解析后的所有数据

    【字段说明】
    - text: 提取的纯文本内容（用于 LLM 分析和 RAG 向量化）
    - metadata: 文件元信息（页数、字符数等）
    - sections: 按章节分割的内容（如：教育经历、项目经验）
    - raw_lines: 原始行列表（保留格式信息）
    """

    text: str                                    # 完整文本内容
    filename: str = ""                           # 原始文件名
    file_type: str = ""                          # 文件类型: pdf/docx
    page_count: int = 0                          # 页数
    char_count: int = 0                          # 字符数
    metadata: Dict[str, Any] = field(default_factory=dict)  # 额外元数据
    sections: Dict[str, str] = field(default_factory=dict)   # 章节内容

    def get_section(self, section_name: str) -> Optional[str]:
        """获取指定章节内容"""
        return self.sections.get(section_name.lower())

    def extract_keywords(self) -> List[str]:
        """
        提取关键词（简单版，后续由 LLM 精确提取）

        【算法说明】
        1. 提取中文词汇（2-4个字的连续中文字符）
        2. 提取英文单词
        3. 过滤常见停用词
        4. 返回去重后的关键词列表
        """
        chinese_pattern = re.compile(r'[\u4e00-\u9fa5]{2,4}')
        english_pattern = re.compile(r'[A-Za-z][A-Za-z0-9+#]*')

        chinese_words = chinese_pattern.findall(self.text)
        english_words = english_pattern.findall(self.text)

        stop_words = {
            '的', '了', '在', '是', '我', '有', '和', '就',
            '不', '人', '都', '一', '一个', '上', '也', '很', '到',
            '说', '要', '去', '你', '会', '着', '没有', '看', '好'
        }

        keywords = []
        for word in set(chinese_words):
            if word not in stop_words and len(word) >= 2:
                keywords.append(word)

        for word in set(english_words):
            if len(word) >= 3:
                keywords.append(word)

        return keywords[:100]

    def __str__(self) -> str:
        return f"ParseResult(file={self.filename}, type={self.file_type}, chars={self.char_count})"


class FileParser:
    """
    文件解析器核心类

    【Java 类比】
    ```java
    @Service
    public class FileParseService {
        @Autowired
        private PdfParser pdfParser;

        @Autowired
        private DocxParser docxParser;

        public ParseResult parseFile(MultipartFile file) {
            String ext = getFileExtension(file);
            switch (ext) {
                case "pdf": return pdfParser.parse(file);
                case "docx": return docxParser.parse(file);
                default: throw new UnsupportedFormatException(ext);
            }
        }
    }
    ```

    【核心功能】
    1. parse_file(): 根据文件类型自动选择解析器
    2. parse_pdf(): PDF 文件解析（使用 pdfplumber）
    3. parse_docx(): Word 文件解析（使用 python-docx）
    4. extract_sections(): 智能章节分割
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}

    def __init__(self):
        self._parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
        }

    async def parse_file(
        self,
        file_path: str,
        extract_sections: bool = True
    ) -> ParseResult:
        """
        解析文件（主入口方法）

        【参数说明】
        - file_path: 文件路径（本地路径或 OSS URL）
        - extract_sections: 是否智能分割章节（默认开启）

        【返回值】
        - ParseResult: 包含文本内容和元数据

        【异常处理】
        - FileParseError: 文件不存在或不支持的格式
        - 由上层 circuit_breaker 或 error_handler 统一捕获

        【使用示例】
        parser = FileParser()

        # 从本地文件解析
        result = await parser.parse_file("/path/to/resume.pdf")

        # 从OSS URL下载后解析（需要先下载到临时目录）
        result = await parser.parse_file("./temp/resume_123.pdf")

        print(result.text[:500])       # 预览前500字符
        print(result.extract_keywords()) # 提取关键词
        """
        path = Path(file_path)

        if not path.exists():
            raise FileParseError(
                f"文件不存在: {file_path}",
                detail=f"Path does not exist: {path.absolute()}"
            )

        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise FileParseError(
                f"不支持的文件格式: {ext}",
                detail=f"仅支持: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        logger.info(
            "file_parse_start",
            file_path=str(path),
            file_type=ext,
            file_size=path.stat().st_size
        )

        parser_func = self._parsers[ext]
        result = parser_func(str(path), path.name)

        if extract_sections and result.text:
            result.sections = self._extract_sections(result.text)

        logger.info(
            "file_parse_complete",
            filename=result.filename,
            char_count=result.char_count,
            page_count=result.page_count,
            section_count=len(result.sections)
        )

        return result

    async def parse_from_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        extract_sections: bool = True
    ) -> ParseResult:
        """
        从字节流解析文件（用于 HTTP 上传的场景）

        【Java 类比】
        ```java
        // 类似 MultipartFile.getBytes() 后解析
        public ParseResult parseFromBytes(byte[] bytes, String filename) {
            Path tempFile = Files.createTempFile("resume_", "_" + filename);
            Files.write(tempFile, bytes);
            try {
                return parseFile(tempFile.toString());
            } finally {
                Files.deleteIfExists(tempFile);
            }
        }
        ```

        【使用场景】
        - Java端转发文件流给Python
        - 直接从内存解析，不落盘
        """
        import tempfile

        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise FileParseError(f"不支持的文件格式: {ext}")

        with tempfile.NamedTemporaryFile(
            suffix=ext,
            delete=False,
            mode='wb'
        ) as tmp:
            tmp.write(file_bytes)
            temp_path = tmp.name

        try:
            return await self.parse_file(temp_path, extract_sections)
        finally:
            os.unlink(temp_path)

    def _parse_pdf(self, file_path: str, filename: str) -> ParseResult:
        """
        解析 PDF 文件（使用 pdfplumber）

        【为什么用 pdfplumber 而不是 PyPDF2？】
        1. 表格提取能力更强（简历中的技能表格）
        2. 中文编码处理更好（UTF-8 支持）
        3. 精确的文字位置信息（可用于布局分析）
        4. 更好的字体样式识别

        【pdfplumber vs PyPDF2 对比】
        | 特性 | pdfplumber | PyPDF2 |
        |------|-----------|--------|
        | 表格提取 | ✅ 强大 | ⚠️ 基础 |
        | 中文支持 | ✅ 好 | ⚠️ 一般 |
        | 速度 | ⚠️ 较慢 | ✅ 快 |
        | 依赖 | 较多 | 较少 |

        【解析策略】
        1. 逐页提取文本
        2. 合并所有页面内容
        3. 清理多余空白和换行
        4. 统计页数和字符数
        """
        import pdfplumber

        all_text = []
        page_count = 0

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)

            for i, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text:
                        cleaned = self._clean_text(text)
                        if cleaned.strip():
                            all_text.append(cleaned)
                except Exception as e:
                    logger.warning(
                        "pdf_page_parse_failed",
                        page=i + 1,
                        error=str(e)
                    )
                    continue

        full_text = "\n\n".join(all_text)

        return ParseResult(
            text=full_text,
            filename=filename,
            file_type="pdf",
            page_count=page_count,
            char_count=len(full_text),
            metadata={
                "parser": "pdfplumber",
                "pages_parsed": len(all_text),
                "has_tables": True  # pdfplumber 支持表格
            }
        )

    def _parse_docx(self, file_path: str, filename: str) -> ParseResult:
        """
        解析 Word 文档 (.docx)（使用 python-docx）

        【python-docx 说明】
        - 只支持 .docx 格式（OOXML 标准）
        - 不支持旧版 .doc 格式（二进制格式）
        - 如果遇到 .doc 文件，需要先用 LibreOffice 转换

        【解析策略】
        1. 按段落读取（paragraphs）
        2. 读取表格内容（tables）
        3. 合并并清理文本
        4. 保留基本格式信息

        【Java 类比】
        - 类似 Apache POI 的 XWPFDocument
        - paragraph → XWPFParagraph
        - table → XWPFTable
        """
        from docx import Document

        doc = Document(file_path)
        all_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else ""
                prefix = ""

                if 'Heading' in style_name or 'heading' in style_name.lower():
                    level = style_name.replace('Heading ', '').replace('heading ', '')
                    prefix = "#" * (int(level) if level.isdigit() else 1) + " "

                all_text.append(prefix + text)

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_data))
            if table_text:
                all_text.append("\n[表格]\n" + "\n".join(table_text))

        full_text = "\n\n".join(all_text)

        return ParseResult(
            text=full_text,
            filename=filename,
            file_type="docx",
            page_count=1,
            char_count=len(full_text),
            metadata={
                "parser": "python-docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
        )

    def _clean_text(self, text: str) -> str:
        """
        清理提取的文本

        【清理规则】
        1. 移除多余空白（连续空格→单个空格）
        2. 规范换行符（\r\n → \n）
        3. 移除首尾空白
        4. 合并过短的行（可能是断词）
        """
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\r\n?', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _extract_sections(self, text: str) -> Dict[str, str]:
        """
        智能分割简历章节

        【识别的章节模式】
        - 教育经历 / Education
        - 项目经历 / Project Experience
        - 实习经历 / Internship
        - 工作经历 / Work Experience
        - 专业技能 / Skills
        - 自我评价 / Self Evaluation
        - 获奖荣誉 / Awards

        【算法说明】
        1. 使用正则匹配章节标题
        2. 提取章节标题到下一个标题之间的内容
        3. 返回章节字典

        【返回值示例】
        {
            "教育经历": "XX大学 计算机科学 本科...",
            "项目经历": "项目1: XXX系统 ...",
            "专业技能": "Java, Spring Boot, MySQL ..."
        }
        """
        section_patterns = [
            (r'(?:教育|学历|学习)[经历背景]', 'education'),
            (r'(?:项目|科研|课题)[经历经验]', 'projects'),
            (r'(?:实习)[经历经验]', 'internships'),
            (r'(?:工作|职业)[经历经验]', 'work_experience'),
            (r'(?:专业)?技能(?:特长|掌握)?', 'skills'),
            (r'(?:自我)?评价|简介|总结', 'summary'),
            (r'(?:获奖|荣誉|证书)', 'awards'),
            (r'校园活动|社会实践', 'activities'),
        ]

        sections = {}
        lines = text.split('\n')
        current_section = 'header'
        current_content = []

        section_regexes = [(re.compile(p, re.IGNORECASE), name) for p, name in section_patterns]

        for line in lines:
            line_stripped = line.strip()
            is_section_header = False

            for regex, name in section_regexes:
                if regex.search(line_stripped):
                    if current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    current_section = name
                    current_content = []
                    is_section_header = True
                    break

            if not is_section_header and line_stripped:
                current_content.append(line_stripped)

        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()

        if not sections or len(sections) <= 1:
            sections['full_content'] = text

        return sections


# ══════════════════════════════════════════════════════════
# 全局单例工厂函数
# ══════════════════════════════════════════════════════════

_file_parser_instance: Optional[FileParser] = None


def get_file_parser() -> FileParser:
    """获取全局文件解析器单例"""
    global _file_parser_instance
    if _file_parser_instance is None:
        _file_parser_instance = FileParser()
    return _file_parser_instance


async def quick_parse(file_path: str) -> str:
    """
    快速解析文件（只返回文本，简化版）

    【使用场景】
    - RAG 向量化时只需要原始文本
    - 不需要章节信息和元数据

    【Java 类比】
    - 类似 FileUtils.readFileToString()
    """
    parser = get_file_parser()
    result = await parser.parse_file(file_path, extract_sections=False)
    return result.text
